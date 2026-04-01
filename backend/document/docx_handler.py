import io
from docx import Document as DocxDocument

_WPS_TXBX = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
_W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"


def fill_docx_template(template_bytes: bytes, placeholders: dict[str, str]) -> bytes:
    """
    Replace {{placeholder}} occurrences in a DOCX template.
    Returns the filled DOCX as bytes.
    """
    doc = DocxDocument(io.BytesIO(template_bytes))

    for para in doc.paragraphs:
        _replace_in_paragraph(para, placeholders)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, placeholders)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, placeholders)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, placeholders)

    from docx.text.paragraph import Paragraph
    for shape in doc.element.body.iter(_WPS_TXBX):
        for para_elem in shape.iter(_W_P):
            para = Paragraph(para_elem, doc)
            _replace_in_paragraph(para, placeholders)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _replace_in_paragraph(paragraph, placeholders: dict[str, str]) -> None:
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    new_text = full_text
    for key, value in placeholders.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", value or "")

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
