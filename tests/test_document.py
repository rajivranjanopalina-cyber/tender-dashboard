# tests/test_document.py
import io
import os
import pytest
from docx import Document as DocxDocument
from unittest.mock import MagicMock, patch


def make_docx_with_placeholders(text: str) -> bytes:
    doc = DocxDocument()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_replaces_placeholder_in_paragraph():
    from backend.document.docx_handler import fill_docx_template
    content = make_docx_with_placeholders("Dear Team, re: {{tender_title}} by {{tender_deadline}}")
    placeholders = {"tender_title": "IT Infrastructure", "tender_deadline": "2026-04-30"}
    result_bytes = fill_docx_template(content, placeholders)
    doc = DocxDocument(io.BytesIO(result_bytes))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "IT Infrastructure" in full_text
    assert "2026-04-30" in full_text
    assert "{{" not in full_text


def test_docx_replaces_placeholder_in_table():
    from backend.document.docx_handler import fill_docx_template
    doc = DocxDocument()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Tender"
    table.cell(0, 1).text = "{{tender_title}}"
    buf = io.BytesIO()
    doc.save(buf)
    placeholders = {"tender_title": "Cloud Migration"}
    result_bytes = fill_docx_template(buf.getvalue(), placeholders)
    doc2 = DocxDocument(io.BytesIO(result_bytes))
    cell_text = doc2.tables[0].cell(0, 1).text
    assert cell_text == "Cloud Migration"


def test_docx_ignores_missing_placeholders():
    from backend.document.docx_handler import fill_docx_template
    content = make_docx_with_placeholders("Value: {{tender_estimated_value}}")
    result_bytes = fill_docx_template(content, {})  # No placeholders provided
    doc = DocxDocument(io.BytesIO(result_bytes))
    # Unreplaced placeholders left as-is
    assert "{{tender_estimated_value}}" in doc.paragraphs[0].text


def test_generator_creates_docx_proposal():
    from backend.document.generator import generate_proposal
    from backend import models

    # Create a minimal DOCX template as bytes
    content = make_docx_with_placeholders("Tender: {{tender_title}}, Deadline: {{tender_deadline}}")

    # Mock ORM objects
    tender = MagicMock(spec=models.Tender)
    tender.id = 1
    tender.title = "IT Networking"
    tender.description = "Upgrade"
    tender.deadline = "2026-04-30"
    tender.published_date = "2026-03-01"
    tender.estimated_value = "50 Lakh"
    tender.source_url = "https://test.gov.in/t/1"
    tender.portal.name = "Test Portal"
    tender.portal.url = "https://test.gov.in"

    template = MagicMock(spec=models.Template)
    template.id = 1
    template.blob_url = "https://blob.vercel-storage.com/templates/template.docx"
    template.file_type = "docx"

    fake_blob_url = "https://blob.vercel-storage.com/proposals/proposal_1_1.docx"

    with patch("backend.document.generator.download_blob", return_value=content) as mock_dl, \
         patch("backend.document.generator.upload_blob", return_value=fake_blob_url) as mock_ul:
        result_url = generate_proposal(tender=tender, template=template)

    assert result_url == fake_blob_url
    mock_dl.assert_called_once_with(template.blob_url)
    # Verify upload was called with the right filename pattern
    upload_call_args = mock_ul.call_args
    assert "proposals/proposal_1_1.docx" in upload_call_args[0][1]
